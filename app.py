import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
import datetime
import io
import random

# --- KONFIGURASI TAMPILAN ---
st.set_page_config(page_title="Generator RPP Cerdas (Offline)", page_icon="🧠", layout="wide")

# --- DATA PENDIDIKAN (OTAK APLIKASI) ---
# Ini adalah database pengetahuan yang saya tanamkan untuk aplikasi ini

LIST_AKHLAK = ["Beriman dan Bertakwa kepada Tuhan YME", "Berakhlak Mulia", "Bernalar Kritis", "Kreatif", "Mandiri", "Bergotong Royong", "Berkebinekaan Global"]
LIST_MODEL = ["Project Based Learning (PjBL)", "Problem Based Learning (PBL)", "Discovery Learning", "Inquiry Learning", "Direct Instruction"]
LIST_METODE = ["Diskusi Kelompok", "Tanya Jawab", "Penugasan", "Simulasi", "Demonstrasi", "Eksperimen", "Studi Kasus"]
LIST_MEDIA = ["Video Pembelajaran", "Slide Presentasi", "Buku Paket", "Lembar Kerja (LKPD)", "Internet", "LCD Proyektor", "Papan Tulis"]
LIST_VERB = ["Menganalisis", "Mengidentifikasi", "Menjelaskan", "Menerapkan", "Mengevaluasi", "Menciptakan", "Membandingkan", "Menyajikan"]

# --- FUNGSI LOGIKA PEMBUAT RPP (EXPERT SYSTEM) ---
def smart_generator(topik, kelas, fase, mapel):
    """
    Fungsi ini mensimulasikan AI dengan menggabungkan aturan pedagogis.
    Ia akan memilih kata yang tepat secara acak namun logis.
    """
    
    # 1. Pilih Komponen Acak yang Relevan
    profil_terpilih = random.sample(LIST_AKHLAK, 2)
    model_terpilih = random.choice(LIST_MODEL)
    metode_terpilih = random.sample(LIST_METODE, 2)
    media_terpilih = random.sample(LIST_MEDIA, 3)
    verb_utama = random.choice(LIST_VERB)
    verb_kedua = random.choice([v for v in LIST_VERB if v != verb_utama])

    # 2. Menyusun Kalimat Dinamis (Template Based Logic)
    
    # Tujuan Pembelajaran (TP)
    tp1 = f"Pertemuan 1: Melalui kegiatan {model_terpilih}, peserta didik mampu {verb_utama} konsep {topik} dengan tepat dan percaya diri."
    tp2 = f"Pertemuan 2: Peserta didik mampu {verb_kedua} implementasi {topik} dalam kehidupan sehari-hari melalui diskusi kelompok."
    
    # Analisis Peserta Didik
    analisis_pd = f"Peserta didik kelas {kelas} memiliki karakteristik aktif dan kritis. Pengetahuan awal tentang {topik} masih terbatas pada pengetahuan umum, sehingga diperlukan pendekatan kontekstual melalui media {media_terpilih[0]}. Sebagian siswa merupakan tipe visual learner."

    # Materi Pelajaran
    materi = f"Jenis Pengetahuan: Konseptual dan Prosedural.\nRelevansi: Materi {topik} sangat relevan dengan tantangan kehidupan sosial remaja saat ini.\nTingkat Kesulitan: Sedang, memerlukan abstraksi tingkat tinggi.\nIntegrasi Karakter: {profil_terpilih[0]} dan {profil_terpilih[1]}."

    # Pertanyaan Pemantik
    pm1 = f"Apakah kamu pernah mengamati fenomena {topik} di lingkungan sekitar?"
    pm2 = f"Bagaimana pendapatmu jika {topik} tidak diterapkan dalam kehidupan bermasyarakat?"
    
    # Kegiatan Inti (Otomatis Disusun)
    kegiatan_inti = f"""
    A. Memahami:
    - Peserta didik mengamati {media_terpilih[0]} terkait {topik}.
    - Peserta didik mengidentifikasi poin-poin penting melalui {metode_terpilih[0]}.

    B. Mengaplikasi:
    - Peserta didik dibagi dalam kelompok untuk menganalisis studi kasus {topik}.
    - Peserta didik membuat {media_terpilih[1]} sebagai hasil analisis.
    - Diferensiasi: Siswa yang visual membuat poster, siswa yang kinestetik melakukan demonstrasi/percobaan.

    C. Merefleksi:
    - Peserta didik mempresentasikan hasil kerja.
    - Peserta didik memberikan saling masukan (peer assessment).
    - Guru mengajak siswa menyimpulkan hikmah pembelajaran {topik} hari ini.
    """

    # Asesmen
    diag_teknik = "Tes Diagnostik Tertulis (Pilihan Ganda/Isian Singkat)"
    form_teknik = f"Observasi sikap menggunakan lembar checklist dan Penilaian Kinerja presentasi kelompok."
    sum_teknik = f"Proyek pembuatan {media_terpilih[1]} tentang {topik} dan Tes Tertulis Akhir."

    # Susun Data Final
    data = {
        "satuan_pendidikan": "SMKN IV SPP-SPMA Singkawang",
        "nama_guru": "Daniel, S.Pd.K",
        "mata_pelajaran": mapel,
        "kelas": kelas,
        "semester": "Ganjil",
        "fase": fase,
        "elemen_pokok": topik,
        "alokasi_waktu": "3 X 3 JP",
        
        "t1_peserta_didik": analisis_pd,
        "t1_materi_pelajaran": materi,
        "t1_profil_lulusan": f"1. {profil_terpilih[0]}\n2. {profil_terpilih[1]}",
        "t1_pertanyaan_pemantik": f"1. {pm1}\n2. {pm2}",
        "t1_sarana": f"Fisik: Ruang kelas, {media_terpilih[2]}.\nDigital: Aplikasi Canva, Video Pembelajaran.",
        
        "t2_cp": f"Memahami dan menerapkan konsep {topik} dalam kehidupan sehari-hari.",
        "t2_tp": f"{tp1}\n{tp2}",
        "t2_pemahaman_bermakna": f"Memahami bahwa {topik} merupakan bagian penting dari pengembangan diri dan masyarakat, serta mampu mengambil sikap bijak terkait isu tersebut.",
        "t2_lintas_disiplin": "PPKn (Nilai Pancasila), Bahasa Indonesia (Komunikasi Efektif)",
        "t2_topik": topik,
        "t2_pedagogis": f"Model: {model_terpilih}\nMetode: {', '.join(metode_terpilih)}",
        "t2_kemitraan": "Orang tua (pendamping di rumah), Tokoh Masyarakat (Narasumber jika memungkinkan)",
        "t2_lingkungan": "Penataan kelas model kelompok (Cluster Seating) untuk memudahkan diskusi.",
        "t2_digital": "Youtube (Sumber Materi), Canva (Media Presentasi), Google Classroom.",
        
        "t3_awal": f"1. Guru membuka pelajaran dengan salam dan doa.\n2. Apersepsi: Mengaitkan materi {topik} dengan pengalaman siswa.\n3. Guru menyampaikan tujuan pembelajaran dan manfaatnya.",
        "t3_awal_prinsip": "Menggembirakan dan Kontekstual.",
        "t3_inti": kegiatan_inti,
        "t3_inti_prinsip": "Berkesadaran, Bermakna, dan Interaktif.",
        "t3_penutup": "1. Siswa membuat kesimpulan pembelajaran (Refleksi 3-2-1).\n2. Guru memberikan penguatan (reinforcement).\n3. Informasi materi pertemuan selanjutnya.",
        "t3_penutup_prinsip": "Reflektif dan Memberdayakan.",
        
        "t4_diagnostik": diag_teknik,
        "t4_diagnostik_kriteria": "Kebenaran jawaban dan pemahaman konsep dasar.",
        "t4_formatif": form_teknik,
        "t4_formatif_kriteria": "Keaktifan, sikap kerjasama, dan ketepatan waktu.",
        "t4_sumatif": sum_teknik,
        "t4_sumatif_kriteria": f"Kedalaman analisis {topik}, kreativitas, dan kebenaran konsep.",
        "t4_tindak_lanjut": "Remedial: Bimbingan individu. Pengayaan: Tugas studi kasus lanjutan."
    }
    
    return data

# --- FUNGSI MEMBUAT WORD (Sama seperti sebelumnya) ---
def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_word_doc(data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    title = doc.add_heading('RENCANA PEMBELAJARAN MENDALAM (DEEP LEARNING)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # INFO UMUM
    info_table = doc.add_table(rows=7, cols=2)
    info_data = [
        ("SATUAN PENDIDIKAN", data.get('satuan_pendidikan', '-')),
        ("NAMA GURU", data.get('nama_guru', '-')),
        ("MATA PELAJARAN", data.get('mata_pelajaran', '-')),
        ("KELAS / SEMESTER", f"{data.get('kelas', '-')} / {data.get('semester', '-')}"),
        ("FASE", data.get('fase', '-')),
        ("ELEMEN/MATERI POKOK", data.get('elemen_pokok', '-')),
        ("ALOKASI WAKTU", data.get('alokasi_waktu', '-')),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i].cells
        row[0].text = label
        row[1].text = f": {value}"
        row[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()

    def add_table(doc, title, headers, rows):
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            set_cell_shading(hdr[i], "D9D9D9")
            hdr[i].paragraphs[0].runs[0].bold = True
        
        for row_data in rows:
            row = table.add_row().cells
            for i, text in enumerate(row_data):
                row[i].text = str(text)
                if i == 0 and len(headers) == 2:
                     if row[i].paragraphs[0].runs:
                        row[i].paragraphs[0].runs[0].bold = True

    # TABEL 1
    t1 = [
        ("Peserta Didik", data.get('t1_peserta_didik', '-')),
        ("Materi Pelajaran", data.get('t1_materi_pelajaran', '-')),
        ("Profil Pelajar", data.get('t1_profil_lulusan', '-')),
        ("Pertanyaan Pemantik", data.get('t1_pertanyaan_pemantik', '-')),
        ("Sarana", data.get('t1_sarana', '-')),
    ]
    add_table(doc, "TABEL 1: IDENTIFIKASI", ["Aspek", "Deskripsi"], t1)

    # TABEL 2
    t2 = [
        ("CP", data.get('t2_cp', '-')),
        ("TP", data.get('t2_tp', '-')),
        ("Pemahaman Bermakna", data.get('t2_pemahaman_bermakna', '-')),
        ("Lintas Disiplin", data.get('t2_lintas_disiplin', '-')),
        ("Topik", data.get('t2_topik', '-')),
        ("Pedagogis", data.get('t2_pedagogis', '-')),
        ("Kemitraan", data.get('t2_kemitraan', '-')),
        ("Lingkungan", data.get('t2_lingkungan', '-')),
        ("Digital", data.get('t2_digital', '-')),
    ]
    add_table(doc, "TABEL 2: DESAIN", ["Komponen", "Rumusan"], t2)

    # TABEL 3
    t3 = [
        ("Awal", data.get('t3_awal', '-'), data.get('t3_awal_prinsip', '-')),
        ("Inti", data.get('t3_inti', '-'), data.get('t3_inti_prinsip', '-')),
        ("Penutup", data.get('t3_penutup', '-'), data.get('t3_penutup_prinsip', '-')),
    ]
    add_table(doc, "TABEL 3: KEGIATAN", ["Tahap", "Kegiatan", "Prinsip"], t3)

    # TABEL 4
    t4 = [
        ("Diagnostik", data.get('t4_diagnostik', '-'), data.get('t4_diagnostik_kriteria', '-')),
        ("Formatif", data.get('t4_formatif', '-'), data.get('t4_formatif_kriteria', '-')),
        ("Sumatif", data.get('t4_sumatif', '-'), data.get('t4_sumatif_kriteria', '-')),
        ("Tindak Lanjut", data.get('t4_tindak_lanjut', '-'), "Sesuai Hasil"),
    ]
    add_table(doc, "TABEL 4: ASESMEN", ["Jenis", "Teknik", "Kriteria"], t4)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- TAMPILAN UTAMA ---
st.title("🧠 Generator RPP Cerdas (Offline Mode)")
st.markdown("Aplikasi ini berjalan **tanpa Internet & Tanpa API**. Logika kurikulum tertanam di dalam sistem.")
st.success("✅ Keuntungan: Gratis Selamanya, Super Cepat, Tidak Perlu Setting API Key.")

col1, col2 = st.columns(2)
with col1:
    mapel = st.text_input("Mata Pelajaran", "Pendidikan Agama Kristen")
    kelas = st.selectbox("Kelas", ["X", "XI", "XII"])
with col2:
    fase = st.selectbox("Fase", ["E", "F", "G"])
    topik = st.text_input("Topik/Materi Pokok", placeholder="Contoh: Ibadah dalam Kehidupan Sehari-hari")

if st.button("⚡ GENERATE RPP SEKARANG", type="primary"):
    if not topik:
        st.warning("Mohon isi Topik terlebih dahulu.")
    else:
        # Proses
        data_hasil = smart_generator(topik, kelas, fase, mapel)
        st.session_state['data_rpp'] = data_hasil
        st.success("✅ RPP Berhasil Dibuat secara Otomatis!")

if 'data_rpp' in st.session_state:
    st.subheader("📥 Download Dokumen")
    data_final = st.session_state['data_rpp']
    
    doc_buffer = create_word_doc(data_final)
    st.download_button(
        label="📄 Download File Word (.docx)",
        data=doc_buffer,
        file_name=f"RPP_{data_final.get('elemen_pokok', 'hasil')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    with st.expander("👀 Lihat Preview Data"):
        st.json(data_final)
